import docx

try:
    print("尝试导入docx模块...")
    doc = docx.Document('3.62 MCP快速入门实战.docx')
    print("成功打开文档！")
    print(f"文档中的段落数量: {len(doc.paragraphs)}")
    print("第一个段落的内容:")
    print(doc.paragraphs[0].text)
except Exception as e:
    print(f"发生错误: {str(e)}")
    print(f"错误类型: {type(e).__name__}")
    import traceback
    traceback.print_exc() 